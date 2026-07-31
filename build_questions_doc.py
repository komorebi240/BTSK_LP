import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DATA_PATH = ROOT / 'questions-data.json'
OUTPUT_PATH = ROOT / 'ty-health-question-bank.docx'

GROUP_NAMES = {
    'female_young': 'Nữ 18–40 tuổi',
    'female_mid': 'Nữ 40 tuổi trở lên',
    'male_young': 'Nam 18–40 tuổi',
    'male_mid': 'Nam 40 tuổi trở lên',
    'child_under2': 'Trẻ dưới 2 tuổi',
    'child_2_5': 'Trẻ 2–5 tuổi',
    'child_6_12': 'Trẻ 6–12 tuổi',
    'child_13_plus': 'Trẻ từ 13 tuổi',
}

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 105, 112)

def set_run_font(run, size=11, bold=None, color=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def add_text(cell, text, bold=False, color=None, size=9.5):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p

def score_text(scores, labels):
    if not scores:
        return '—'
    return '; '.join(f"{labels.get(axis, axis)}: {value}" for axis, value in scores.items())

def add_question_table(doc, questions, labels):
    # No. / Question / Choices & score / Selection rule
    table = doc.add_table(rows=1, cols=4)
    widths = [420, 2780, 4620, 1540]
    headers = ['#', 'Câu hỏi', 'Đáp án và điểm theo trục', 'Cách chọn']
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, 'E8EEF5')
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)

    for index, question in enumerate(questions, 1):
        row = table.add_row()
        cells = row.cells
        add_text(cells[0], str(index), bold=True, color=DARK_BLUE)
        add_text(cells[1], question.get('tag', 'Không phân loại'), bold=True, color=DARK_BLUE)
        add_text(cells[1], question.get('text', ''), size=9.5)
        if question.get('sub'):
            add_text(cells[1], question['sub'], color=MUTED, size=8.8)

        for opt_index, option in enumerate(question.get('opts', []), 1):
            prefix = f"{opt_index}. "
            if option.get('exclusive'):
                prefix += '[Chọn riêng] '
            add_text(cells[2], prefix + option.get('text', ''), bold=True, size=9.0)
            if option.get('sub'):
                add_text(cells[2], option['sub'], color=MUTED, size=8.5)
            add_text(cells[2], 'Điểm: ' + score_text(option.get('scores', {}), labels), color=MUTED, size=8.5)

        if question.get('multi'):
            add_text(cells[3], f"Chọn nhiều, tối đa {question.get('maxSelect', 3)} ý", bold=True, size=9.0)
            add_text(cells[3], question.get('multiNote', 'Đáp án “chọn riêng” không đi kèm đáp án khác.'), color=MUTED, size=8.5)
        else:
            add_text(cells[3], 'Chọn 1 đáp án', bold=True, size=9.0)
        add_text(cells[3], 'Key: ' + question.get('key', '—'), color=MUTED, size=8.3)

    set_table_geometry(table, widths)
    # Repeat the header row across page breaks.
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [('Heading 1', 16, BLUE, 14, 7), ('Heading 2', 13, BLUE, 10, 5), ('Heading 3', 12, DARK_BLUE, 8, 4)]:
        style = doc.styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Trang ')
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement('w:fldChar'); fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar'); fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1); run._r.append(instr_text); run._r.append(fld_char2)

def main():
    data = json.loads(DATA_PATH.read_text(encoding='utf-8-sig'))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    setup_styles(doc)

    header = section.header.paragraphs[0]
    header.text = 'TY HEALTH SUPPLEMENTS | NGÂN HÀNG CÂU HỎI'
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs: set_run_font(run, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run('Ngân hàng câu hỏi và thang điểm')
    set_run_font(title_run, size=24, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph('TY Health Check | Trích xuất từ index.html')
    subtitle.paragraph_format.space_after = Pt(14)
    for run in subtitle.runs: set_run_font(run, size=11, color=MUTED)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    r = intro.add_run('Phạm vi: '); set_run_font(r, bold=True, color=DARK_BLUE)
    r = intro.add_run('Toàn bộ câu hỏi theo hồ sơ, các lựa chọn và điểm cộng cho từng trục sức khỏe.'); set_run_font(r)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    r = note.add_run('Quy tắc câu hỏi nhiều đáp án: '); set_run_font(r, bold=True, color=DARK_BLUE)
    r = note.add_run('chọn tối đa 3 ý. Lựa chọn “chọn riêng” không được kết hợp với các lựa chọn khác.'); set_run_font(r)

    doc.add_heading('Trục điểm', level=1)
    axes = doc.add_table(rows=1, cols=2)
    for cell, header_text in zip(axes.rows[0].cells, ['Mã trục', 'Ý nghĩa']):
        set_cell_shading(cell, 'E8EEF5')
        run = cell.paragraphs[0].add_run(header_text); set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
    for code, label in data['AXIS_LABELS'].items():
        row = axes.add_row().cells
        add_text(row[0], code, bold=True, color=DARK_BLUE)
        add_text(row[1], label)
    set_table_geometry(axes, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.add_heading('Câu hỏi khởi tạo', level=1)
    add_question_table(doc, [data['profileQuestion'], data['childAgeQuestion']], data['AXIS_LABELS'])

    for group_key, questions in data['QUESTION_SETS'].items():
        doc.add_heading(GROUP_NAMES.get(group_key, group_key), level=1)
        add_question_table(doc, questions, data['AXIS_LABELS'])

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)

if __name__ == '__main__':
    main()
